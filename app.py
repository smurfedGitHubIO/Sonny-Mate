import docx
from docx.shared import Inches

doc = docx.Document()

# doc.add_paragraph('Hello World!')
# paraObj1 = doc.add_paragraph('Second paragraph')
# paraObj2 = doc.add_paragraph('Third paragraph')
# paraObj1.add_run('This text is being added to the second paragraph')
# doc.add_page_break()

# hard coded
for i in range(1,4):
    doc.add_picture('./images/image' + str(i) + '.jpg', width=Inches(3.0))
    doc.add_page_break()

doc.save('./tests/sample.docx')
